# -*- coding: utf-8 -*-
"""
Fill external lecturer payment receipts (.docx) with personal data and images.

Usage:
    python fill_receipt.py <input.docx> [input2.docx ...] [--output-dir <dir>]

The script:
  1. Reads personal data from personal-data-filler skill
  2. Identifies form fields by keyword matching
  3. Fills text fields (name, ID, address, unit, phone, bank info)
  4. Inserts signature, ID card, and bank passbook images
  5. Outputs as "<original>_filled.docx" (never overwrites originals)
"""
import argparse
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


# --- Configuration ---

# Personal data JSON path
PERSONAL_DATA_PATH = Path(r"C:\Users\user\.claude\skills\personal-data-filler\references\personal-data.json")

# Image paths (checked at runtime; missing images are skipped)
IMAGE_PATHS = {
    "signature": Path(r"C:\Users\user\Pictures\曾慶良簽名.jpg"),
    "id_back": Path(r"C:\Users\user\Pictures\身分證背面.jpg"),
    "bank": Path(r"C:\Users\user\Pictures\帳戶影本.jpg"),
}

# Default image sizes (cm)
SIGNATURE_WIDTH_CM = 2.5
ID_CARD_WIDTH_CM = 8.5
BANK_IMAGE_HEIGHT_CM = 5.5
BANK_IMAGE_WIDTH_CM = 8.0


def load_personal_data():
    """Load personal data from JSON file."""
    with open(PERSONAL_DATA_PATH, "r", encoding="utf-8") as f:
        raw = json.load(f)

    # Parse bank name and branch from combined field
    bank_full = raw.get("bank_name", "")
    parts = bank_full.split()
    bank_name = parts[0].replace("銀行", "") if parts else ""
    branch = parts[1].replace("分行", "") if len(parts) > 1 else ""

    return {
        "name": raw.get("name", ""),
        "id_number": raw.get("id_number", ""),
        "address": raw.get("address", ""),
        "unit": raw.get("service_unit", ""),
        "phone": raw.get("phone", ""),
        "bank_name": bank_name,
        "branch": branch,
        "bank_code": raw.get("bank_code", ""),
        "account": raw.get("account_number", ""),
    }


def get_para_text(para):
    """Get combined text from all runs in a paragraph."""
    return "".join(run.text for run in para.runs)


def get_image_orientation(img_path):
    """Return 'portrait' or 'landscape' based on image dimensions."""
    with Image.open(img_path) as img:
        w, h = img.size
    return "portrait" if h > w else "landscape"


def collect_unique_paragraphs(doc):
    """
    Collect all paragraphs from the document, deduplicating merged cells.
    Returns list of paragraphs.
    """
    seen_ids = set()
    paragraphs = []

    # Table paragraphs first (most form content is in tables)
    for table in doc.tables:
        for row in table.rows:
            seen_cells = set()
            for cell in row.cells:
                if id(cell) in seen_cells:
                    continue
                seen_cells.add(id(cell))
                for p in cell.paragraphs:
                    if id(p) not in seen_ids:
                        seen_ids.add(id(p))
                        paragraphs.append(p)

    # Body-level paragraphs
    for p in doc.paragraphs:
        if id(p) not in seen_ids:
            seen_ids.add(id(p))
            paragraphs.append(p)

    return paragraphs


def fill_text_fields(paragraphs, data):
    """Fill text fields in paragraphs. Returns set of filled field names."""
    filled = set()

    for para in paragraphs:
        text = get_para_text(para)

        # --- Name + Signature ---
        if "姓名：" in text and "簽名或蓋章" in text and "name" not in filled:
            filled.add("name")
            for i, run in enumerate(para.runs):
                if "姓名：" in run.text:
                    run.text = f"姓名：{data['name']} "
                    break

            # Insert signature image right after name
            sig_path = IMAGE_PATHS.get("signature")
            if sig_path and sig_path.exists():
                name_elem = None
                for run in para.runs:
                    if data["name"] in run.text:
                        name_elem = run._element
                        break
                if name_elem is not None:
                    new_run = para.add_run()
                    new_run.add_picture(str(sig_path), width=Cm(SIGNATURE_WIDTH_CM))
                    name_elem.addnext(new_run._element)

        # --- ID Number ---
        elif "身分證字號：" in text and "id" not in filled:
            filled.add("id")
            found = False
            for run in para.runs:
                if "身分證字號：" in run.text:
                    run.text = f"身分證字號：{data['id_number']}"
                    found = True
                    continue
                if found and run.text.strip() == "":
                    run.text = ""

        # --- Address ---
        elif "戶籍地址" in text and "含里" in text and "addr" not in filled:
            filled.add("addr")
            for run in para.runs:
                if run.text.strip() == "：":
                    run.text = f"：{data['address']}"
                    break

        # --- Service Unit + Phone ---
        elif "服務" in text and "連絡電話" in text and "unit" not in filled:
            filled.add("unit")
            found_label = False
            for run in para.runs:
                if "單位：" in run.text:
                    found_label = True
                    continue
                if found_label and run.text.strip() == "" and "連絡" not in run.text:
                    run.text = data["unit"]
                    found_label = False
                    continue
                if "連絡電話：" in run.text:
                    run.text = f"連絡電話：{data['phone']}"

        # --- Bank Info ---
        elif "銀行" in text and "郵局" in text and "分行" in text and "bank" not in filled:
            filled.add("bank")

            for run in para.runs:
                t = run.text.strip()
                if t == "銀行":
                    run.text = f"{data['bank_name']}銀行"
                elif "分行。" in run.text:
                    run.text = f"{data['branch']}分行。"
                elif "銀行代號：" in run.text:
                    run.text = f"銀行代號：{data['bank_code']}"

            # Account number
            found_acct = False
            for run in para.runs:
                if "帳號" in run.text and "匯款" not in run.text:
                    found_acct = True
                    continue
                if found_acct and "：" in run.text:
                    run.text = f"：{data['account']}"
                    found_acct = False
                    break

            # Clear blank runs after bank code
            in_code = False
            for run in para.runs:
                if "銀行代號：" in run.text:
                    in_code = True
                    continue
                if in_code and run.text.strip() == "":
                    run.text = ""
                    in_code = False

    return filled


def insert_images_in_table(doc, filled):
    """Insert ID card and bank passbook images into table cells."""

    # --- ID Card Image ---
    id_path = IMAGE_PATHS.get("id_back")
    if id_path and id_path.exists() and "img_id" not in filled:
        for table in doc.tables:
            seen = set()
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    if id(cell) in seen:
                        continue
                    seen.add(id(cell))
                    if cell.text.strip() == "黏貼處":
                        filled.add("img_id")
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = ""
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(str(id_path), width=Cm(ID_CARD_WIDTH_CM))
                        break

    # --- Bank Account Image ---
    bank_path = IMAGE_PATHS.get("bank")
    if bank_path and bank_path.exists() and "img_bank" not in filled:
        for table in doc.tables:
            found_header = False
            for row_idx, row in enumerate(table.rows):
                row_text = ""
                seen_row = set()
                for cell in row.cells:
                    if id(cell) not in seen_row:
                        seen_row.add(id(cell))
                        row_text += cell.text

                if "匯款資料" in row_text:
                    found_header = True
                    continue

                if found_header and row_text.strip() == "":
                    filled.add("img_bank")
                    cell = row.cells[0]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()

                    # Auto-detect orientation for proper sizing
                    orientation = get_image_orientation(str(bank_path))
                    if orientation == "portrait":
                        run.add_picture(str(bank_path), height=Cm(BANK_IMAGE_HEIGHT_CM))
                    else:
                        run.add_picture(str(bank_path), width=Cm(BANK_IMAGE_WIDTH_CM))
                    break


def fill_receipt(input_path, output_path):
    """Fill a single receipt document."""
    doc = Document(input_path)
    paragraphs = collect_unique_paragraphs(doc)
    data = load_personal_data()

    filled = fill_text_fields(paragraphs, data)
    insert_images_in_table(doc, filled)

    doc.save(output_path)
    return filled


def main():
    parser = argparse.ArgumentParser(description="Fill receipt forms with personal data")
    parser.add_argument("files", nargs="+", help="Input .docx files")
    parser.add_argument("--output-dir", default=None, help="Output directory (default: same as input)")
    args = parser.parse_args()

    for fpath in args.files:
        fpath = os.path.abspath(fpath)
        if not os.path.exists(fpath):
            print(f"[SKIP] File not found: {fpath}")
            continue

        base, ext = os.path.splitext(os.path.basename(fpath))
        out_dir = args.output_dir or os.path.dirname(fpath)
        out_path = os.path.join(out_dir, f"{base}_filled{ext}")

        try:
            filled = fill_receipt(fpath, out_path)
            fields = ", ".join(sorted(filled))
            print(f"[OK] {os.path.basename(fpath)} -> {os.path.basename(out_path)} ({fields})")
        except Exception as e:
            print(f"[ERROR] {os.path.basename(fpath)}: {e}")


if __name__ == "__main__":
    main()
