#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 轉 Word (.docx) 工具
===============================
支援功能：
  1. 標題（H1-H6）
  2. 粗體、斜體、粗斜體、刪除線、行內程式碼
  3. 無序清單、有序清單（含巢狀）
  4. 表格（含標頭）
  5. 程式碼區塊（含語法標示名稱）
  6. 引用區塊（blockquote）
  7. 水平線
  8. 超連結
  9. 圖片（本地檔案）
  10. YAML Front Matter 自動忽略或轉為文件屬性
  11. 自訂樣式模板

依賴套件安裝：
  pip install python-docx markdown beautifulsoup4 pyyaml
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("❌ 請先安裝 python-docx：pip install python-docx")
    sys.exit(1)

try:
    import markdown
except ImportError:
    print("❌ 請先安裝 markdown：pip install markdown")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError:
    print("❌ 請先安裝 beautifulsoup4：pip install beautifulsoup4")
    sys.exit(1)


# ─────────────────────────────────────────────
# YAML Front Matter 處理
# ─────────────────────────────────────────────

def extract_front_matter(md_text: str) -> tuple[dict, str]:
    """從 Markdown 中提取 YAML Front Matter，回傳 (metadata, content)。"""
    pattern = r"^---\s*\n(.*?)\n---\s*\n"
    match = re.match(pattern, md_text, re.DOTALL)
    if not match:
        return {}, md_text

    yaml_str = match.group(1)
    content = md_text[match.end():]

    try:
        import yaml
        metadata = yaml.safe_load(yaml_str) or {}
    except ImportError:
        metadata = {}
    except Exception:
        metadata = {}

    return metadata, content


# ─────────────────────────────────────────────
# 樣式設定
# ─────────────────────────────────────────────

def setup_styles(doc: Document, font_name: str = "微軟正黑體", en_font: str = "Calibri"):
    """設定文件預設樣式。"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = en_font
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # 設定標題樣式
    heading_sizes = {1: 22, 2: 18, 3: 15, 4: 13, 5: 11, 6: 11}
    heading_colors = {
        1: RGBColor(0x1A, 0x1A, 0x2E),
        2: RGBColor(0x2D, 0x2D, 0x44),
        3: RGBColor(0x3D, 0x3D, 0x5C),
        4: RGBColor(0x4A, 0x4A, 0x6A),
        5: RGBColor(0x5A, 0x5A, 0x7A),
        6: RGBColor(0x6A, 0x6A, 0x8A),
    }

    for level in range(1, 7):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.size = Pt(heading_sizes[level])
            h_style.font.color.rgb = heading_colors[level]
            h_style.font.name = en_font
            h_style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # 建立程式碼區塊樣式
    if "Code Block" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
        pf = code_style.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.left_indent = Cm(0.5)

    # 建立引用樣式
    if "Block Quote" not in [s.name for s in doc.styles]:
        quote_style = doc.styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
        quote_style.font.name = en_font
        quote_style.font.size = Pt(11)
        quote_style.font.italic = True
        quote_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        quote_style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        pf = quote_style.paragraph_format
        pf.left_indent = Cm(1.0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)


# ─────────────────────────────────────────────
# HTML → docx 轉換引擎
# ─────────────────────────────────────────────

class MarkdownToDocxConverter:
    """將 Markdown（經 HTML 中介）轉換為 Word 文件。"""

    def __init__(self, doc: Document, md_file_dir: Path = None):
        self.doc = doc
        self.md_file_dir = md_file_dir or Path(".")
        self.list_level = 0

    def convert(self, md_text: str):
        """主要轉換流程。"""
        # Markdown → HTML
        extensions = [
            "tables",
            "fenced_code",
            "codehilite",
            "toc",
            "nl2br",
            "sane_lists",
            "smarty",
        ]
        html = markdown.markdown(md_text, extensions=extensions)
        soup = BeautifulSoup(html, "html.parser")

        # 遞迴處理 HTML 節點
        for element in soup.children:
            self._process_element(element)

    def _process_element(self, element):
        """遞迴處理單一 HTML 元素。"""
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                p = self.doc.add_paragraph(text)
            return

        if not isinstance(element, Tag):
            return

        tag = element.name

        # 標題
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            p = self.doc.add_heading("", level=level)
            self._add_inline(p, element)

        # 段落
        elif tag == "p":
            # 檢查是否只包含圖片
            imgs = element.find_all("img")
            if imgs and len(list(element.children)) == len(imgs):
                for img in imgs:
                    self._add_image(img)
            else:
                p = self.doc.add_paragraph()
                self._add_inline(p, element)

        # 無序清單
        elif tag == "ul":
            self._process_list(element, ordered=False)

        # 有序清單
        elif tag == "ol":
            self._process_list(element, ordered=True)

        # 表格
        elif tag == "table":
            self._process_table(element)

        # 程式碼區塊
        elif tag == "pre":
            code_el = element.find("code")
            code_text = code_el.get_text() if code_el else element.get_text()
            # 取得語言名稱
            lang = ""
            if code_el and code_el.get("class"):
                for cls in code_el["class"]:
                    if cls.startswith("language-"):
                        lang = cls.replace("language-", "")
                        break
            if lang:
                p = self.doc.add_paragraph(f"[{lang}]", style="Code Block")
            for line in code_text.rstrip("\n").split("\n"):
                self.doc.add_paragraph(line, style="Code Block")

        # 引用區塊
        elif tag == "blockquote":
            for child in element.children:
                if isinstance(child, Tag) and child.name == "p":
                    p = self.doc.add_paragraph(style="Block Quote")
                    self._add_inline(p, child)
                elif isinstance(child, Tag):
                    self._process_element(child)

        # 水平線
        elif tag == "hr":
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)

        # div 等容器
        elif tag in ("div", "section", "article", "main"):
            for child in element.children:
                self._process_element(child)

        # 其他未處理的標籤
        else:
            text = element.get_text().strip()
            if text:
                p = self.doc.add_paragraph(text)

    def _add_inline(self, paragraph, element):
        """處理行內格式（粗體、斜體、連結、行內程式碼等）。"""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    paragraph.add_run(text)

            elif isinstance(child, Tag):
                if child.name == "strong" or child.name == "b":
                    run = paragraph.add_run(child.get_text())
                    run.bold = True

                elif child.name == "em" or child.name == "i":
                    run = paragraph.add_run(child.get_text())
                    run.italic = True

                elif child.name == "code":
                    run = paragraph.add_run(child.get_text())
                    run.font.name = "Consolas"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

                elif child.name == "del" or child.name == "s":
                    run = paragraph.add_run(child.get_text())
                    run.font.strike = True

                elif child.name == "a":
                    href = child.get("href", "")
                    text = child.get_text()
                    run = paragraph.add_run(text)
                    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                    run.font.underline = True

                elif child.name == "img":
                    self._add_image(child)

                elif child.name == "br":
                    paragraph.add_run("\n")

                else:
                    # 遞迴處理巢狀行內元素
                    self._add_inline(paragraph, child)

    def _process_list(self, element, ordered=False):
        """處理清單（含巢狀）。"""
        for i, li in enumerate(element.find_all("li", recursive=False)):
            # 判斷是否有巢狀清單
            nested_lists = li.find_all(["ul", "ol"], recursive=False)

            # 加入清單項目文字
            if ordered:
                prefix = f"{i + 1}. "
                style = "List Number"
            else:
                prefix = ""
                style = "List Bullet"

            # 取得純文字（排除巢狀清單的文字）
            text_parts = []
            for child in li.children:
                if isinstance(child, NavigableString):
                    t = str(child).strip()
                    if t:
                        text_parts.append(t)
                elif isinstance(child, Tag) and child.name not in ("ul", "ol"):
                    text_parts.append(child.get_text().strip())

            text = " ".join(text_parts)
            if text:
                try:
                    p = self.doc.add_paragraph(prefix + text, style=style)
                except KeyError:
                    p = self.doc.add_paragraph(prefix + text)

            # 遞迴處理巢狀清單
            for nested in nested_lists:
                self.list_level += 1
                is_ordered = nested.name == "ol"
                self._process_list(nested, ordered=is_ordered)
                self.list_level -= 1

    def _process_table(self, element):
        """處理 HTML 表格 → Word 表格。"""
        rows_data = []

        # 取得表頭
        thead = element.find("thead")
        if thead:
            for tr in thead.find_all("tr"):
                cells = [td.get_text().strip() for td in tr.find_all(["th", "td"])]
                rows_data.append(cells)

        # 取得表身
        tbody = element.find("tbody")
        body_source = tbody if tbody else element
        for tr in body_source.find_all("tr", recursive=False if tbody else True):
            # 跳過已在 thead 處理過的
            if thead and tr.parent == thead:
                continue
            cells = [td.get_text().strip() for td in tr.find_all(["th", "td"])]
            if cells:
                rows_data.append(cells)

        if not rows_data:
            return

        # 計算欄數
        col_count = max(len(row) for row in rows_data)
        # 補齊欄數
        for row in rows_data:
            while len(row) < col_count:
                row.append("")

        # 建立 Word 表格
        table = self.doc.add_table(rows=len(rows_data), cols=col_count)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows_data):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text

                # 表頭粗體
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # 表格後加空行
        self.doc.add_paragraph()

    def _add_image(self, img_tag):
        """插入圖片。"""
        src = img_tag.get("src", "")
        alt = img_tag.get("alt", "")

        if not src:
            return

        # 嘗試解析本地路徑
        img_path = self.md_file_dir / src
        if img_path.exists():
            try:
                self.doc.add_picture(str(img_path), width=Inches(5.5))
                if alt:
                    p = self.doc.add_paragraph(alt)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            except Exception as e:
                self.doc.add_paragraph(f"[圖片載入失敗: {src}] {e}")
        else:
            self.doc.add_paragraph(f"[圖片: {alt or src}]")


# ─────────────────────────────────────────────
# 主要轉換函式
# ─────────────────────────────────────────────

def convert_md_to_docx(
    md_path: Path,
    output_path: Path | None = None,
    template_path: Path | None = None,
    font_name: str = "微軟正黑體",
    en_font: str = "Calibri",
) -> Path:
    """
    將 Markdown 檔案轉換為 Word (.docx)。

    參數：
      md_path       - 輸入 .md 檔案路徑
      output_path   - 輸出 .docx 路徑（預設同目錄同名）
      template_path - Word 模板檔案路徑（.docx）
      font_name     - 中文字型名稱
      en_font       - 英文字型名稱
    """
    if not md_path.exists():
        raise FileNotFoundError(f"找不到檔案：{md_path}")

    if output_path is None:
        output_path = md_path.with_suffix(".docx")

    print(f"\n🔄 轉換：{md_path.name}")

    # 讀取 Markdown
    md_text = md_path.read_text(encoding="utf-8")

    # 提取 Front Matter
    metadata, content = extract_front_matter(md_text)

    # 建立文件（使用模板或空白）
    if template_path and template_path.exists():
        print(f"  📋 使用模板：{template_path.name}")
        doc = Document(str(template_path))
    else:
        doc = Document()

    # 設定樣式
    setup_styles(doc, font_name=font_name, en_font=en_font)

    # 設定文件屬性
    if metadata:
        core = doc.core_properties
        if "title" in metadata:
            core.title = str(metadata["title"])
        if "author" in metadata:
            core.author = str(metadata["author"])
        if "subject" in metadata:
            core.subject = str(metadata["subject"])
        print(f"  📝 Front Matter → 文件屬性：{list(metadata.keys())}")

    # 轉換內容
    converter = MarkdownToDocxConverter(doc, md_file_dir=md_path.parent)
    converter.convert(content)

    # 儲存
    doc.save(str(output_path))
    print(f"  ✅ 已輸出：{output_path}")
    return output_path


def batch_convert(
    folder: Path,
    output_folder: Path | None = None,
    template_path: Path | None = None,
    font_name: str = "微軟正黑體",
    en_font: str = "Calibri",
):
    """批次轉換資料夾內所有 .md 檔案。"""
    md_files = list(folder.glob("*.md"))
    if not md_files:
        print(f"⚠️  在 {folder} 中找不到任何 .md 檔案。")
        return

    out_dir = output_folder or folder
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"📂 批次轉換 {len(md_files)} 個 Markdown，輸出到：{out_dir}")
    success, failed = 0, []

    for md_file in md_files:
        out = out_dir / md_file.with_suffix(".docx").name
        try:
            convert_md_to_docx(
                md_file, out,
                template_path=template_path,
                font_name=font_name,
                en_font=en_font,
            )
            success += 1
        except Exception as e:
            print(f"  ❌ 失敗：{md_file.name}，原因：{e}")
            failed.append(md_file.name)

    print(f"\n✅ 完成：{success}/{len(md_files)} 個成功")
    if failed:
        print("❌ 失敗清單：" + ", ".join(failed))


# ─────────────────────────────────────────────
# CLI 入口
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Markdown 轉 Word (.docx) 工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
範例：
  # 轉換單一 Markdown
  python md_to_docx.py report.md

  # 指定輸出路徑
  python md_to_docx.py report.md -o output.docx

  # 使用 Word 模板
  python md_to_docx.py report.md --template company_template.docx

  # 指定字型
  python md_to_docx.py report.md --font "標楷體" --en-font "Times New Roman"

  # 批次轉換整個資料夾
  python md_to_docx.py ./docs/ --batch

  # 批次轉換，輸出到指定資料夾
  python md_to_docx.py ./docs/ --batch -o ./output/
        """,
    )
    parser.add_argument("input", help=".md 檔案路徑 或 資料夾路徑（搭配 --batch）")
    parser.add_argument("-o", "--output", help="輸出 .docx 路徑 或 輸出資料夾（搭配 --batch）")
    parser.add_argument("--template", help="Word 模板檔案路徑（.docx）")
    parser.add_argument("--font", default="微軟正黑體", help="中文字型（預設：微軟正黑體）")
    parser.add_argument("--en-font", default="Calibri", help="英文字型（預設：Calibri）")
    parser.add_argument("--batch", action="store_true", help="批次轉換資料夾內所有 .md")

    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output) if args.output else None
    template_path = Path(args.template) if args.template else None

    if args.batch or input_path.is_dir():
        if not input_path.is_dir():
            print(f"❌ 批次模式需要指定資料夾，但 {input_path} 不是資料夾。")
            sys.exit(1)
        batch_convert(
            input_path, output_path,
            template_path=template_path,
            font_name=args.font,
            en_font=args.en_font,
        )
    else:
        if not input_path.suffix.lower() == ".md":
            print(f"❌ 輸入檔案必須是 .md 格式：{input_path}")
            sys.exit(1)
        convert_md_to_docx(
            input_path, output_path,
            template_path=template_path,
            font_name=args.font,
            en_font=args.en_font,
        )


if __name__ == "__main__":
    main()
